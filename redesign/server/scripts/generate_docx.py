import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Title
    title = doc.add_paragraph()
    run_title = title.add_run("📦 IPOS5 — Dokumentasi Integrasi & Skema Field Database MongoDB")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    # Connection metadata box
    info_p = doc.add_paragraph()
    info_run = info_p.add_run(
        "Database: ipos5_reporting\n"
        "Server Primary: 192.168.5.219:27017 (Primary MongoDB Remote)\n"
        "Auth: mongodb://Valdric:****@192.168.5.219:27017/ipos5_reporting?authSource=admin\n"
        "Koneksi diatur di: connections.json → dikelola oleh DbConnection.js"
    )
    info_run.font.size = Pt(9.5)
    info_run.font.italic = True
    info_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph() # space

    # Section 1: Daftar Koleksi
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("🗃️ 1. Daftar Koleksi MongoDB (Collections)")
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    t1_data = [
        ("Nama Koleksi", "Deskripsi Utama", "Dipakai Di"),
        ("transaksi", "Data paket / connote kiriman", "Checker, Transaksi, Manifest, Gate, Dashboard"),
        ("master_kendaraan", "Data armada mobil operasional", "Checker, Fleet, Gate Monitoring"),
        ("master_kantor", "Data kantor / nopend / KC-KC", "Checker, Dashboard, Semua lookup kantor"),
        ("detail_route", "Waypoint stop-stop per rute", "Checker, Route Journey, Estimasi"),
        ("master_route_nopen", "Header rute operasional", "Checker, Route Journey, Jadwal, Estimasi"),
        ("route_journeys", "Perjalanan aktif per kendaraan per tanggal", "Checker, Route Journey, Estimasi"),
        ("jadwal_transportasi", "Jadwal pickup/transport harian", "Jadwal Pickup, Jadwal Transportasi, Checker"),
        ("template_jadwal_transportasi", "Template jadwal default per rute", "Jadwal, Template"),
        ("manifests (manifest_master & detail)", "Data manifest kantong kiriman", "Transaksi / Manifest controller"),
        ("tracking_events", "Event log tracking per paket", "Checker (timeline)"),
        ("users", "Data akun pengguna sistem", "Login, Profil, Manajemen User"),
        ("master_produk", "Katalog layanan & produk Pos", "Master Produk")
    ]
    
    t1 = doc.add_table(rows=len(t1_data), cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths1 = [Inches(2.0), Inches(2.3), Inches(2.5)]
    
    for r_idx, row in enumerate(t1.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = widths1[c_idx]
            cell.text = t1_data[r_idx][c_idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if r_idx == 0:
                set_cell_background(cell, "003366")
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(9.5)
            else:
                set_cell_background(cell, "F2F4F7" if r_idx % 2 == 1 else "FFFFFF")
                for r in p.runs:
                    r.font.size = Pt(9.0)

    doc.add_paragraph() # space

    # Section 2: Halaman Web & Koleksi MongoDB Yang Dipakai (Complete 15 pages)
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("🖥️ 2. Halaman Web & Koleksi MongoDB Yang Dipakai")
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    pages_data = [
        {
            "title": "1. 🔐 Halaman Login (/login)",
            "meta": "File: Login.jsx | API: POST /api/auth/login | Controller: AuthController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Diambil"),
                ("users", "findOne({ username })", "Username, password_hash, role, nama, NIP, branch")
            ],
            "note": "Ada fallback user bawaan (admin/admin, sari/sari, operator/operator) jika database offline."
        },
        {
            "title": "2. 📊 Halaman Dashboard (/)",
            "meta": "File: Dashboard.jsx | API: GET /api/dashboard-stats | Controller: TransactionController.js (getDashboardStats)",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("transaksi", "aggregate / count", "Statistik paket: total, in-transit, delivered, entry"),
                ("master_kendaraan", "find({})", "Jumlah armada aktif"),
                ("route_journeys", "find({ status: 'IN_PROGRESS' })", "Jumlah perjalanan aktif hari ini"),
                ("master_kantor", "find({})", "Jumlah kantor terdaftar")
            ]
        },
        {
            "title": "3. 🔍 Halaman Package Tracking / Checker (/checker)",
            "meta": "File: Checker.jsx | API: GET /api/checker/:connoteOrNopol, GET /api/kendaraan | Controllers: TransactionController.js, KendaraanController.js",
            "subsections": [
                {
                    "subtitle": "Saat mencari Nomor Resi / Connote:",
                    "table": [
                        ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan di UI"),
                        ("transaksi", "findOne({ connote_code })", "Berat, layanan, pengirim, penerima, asal, tujuan, status"),
                        ("master_kantor", "find({ nopend: [...] })", "Nama kantor asal & tujuan"),
                        ("master_route_nopen", "find({ nopen_asal })", "Rute yang cocok untuk paket"),
                        ("detail_route", "find({ route_id })", "Waypoints / stop rute"),
                        ("jadwal_transportasi", "findOne({ asal, tujuan, tanggal })", "Jadwal kendaraan yang relevan"),
                        ("template_jadwal_transportasi", "findOne({ ... })", "Fallback template jadwal"),
                        ("master_kendaraan", "findOne({ kendaraan_id })", "Info kendaraan yang dijadwalkan"),
                        ("route_journeys", "findOne({ cargo.connote_code })", "Journey aktif yang memuat paket ini"),
                        ("tracking_events", "find({ connote_code })", "Timeline event tracking (ENTRY → IN_TRANSIT → DELIVERED)")
                    ]
                },
                {
                    "subtitle": "Saat mencari Plat Nomor Kendaraan (mode armada):",
                    "table": [
                        ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan di UI"),
                        ("master_kendaraan", "find({}) → normalisasi nopol", "Profil armada: driver, jenis, kapasitas, home base, rute"),
                        ("route_journeys", "findOne({ vehicle_nopol, journey_date })", "Journey & cargo aktif hari ini"),
                        ("detail_route", "find({ route_id, status: 'AKTIF' })", "Waypoints multi-stop route stepper"),
                        ("master_kantor", "find({ nopend: [...stopCodes] })", "Nama kantor per stop di stepper")
                    ]
                },
                {
                    "subtitle": "List Fleet Mobil Panel (saat klik tab List Fleet):",
                    "table": [
                        ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan di UI"),
                        ("master_kendaraan", "find({}) (via GET /api/kendaraan)", "Kartu armada: nopol, driver, jenis, max_capacity_kg, rute, home_base")
                    ]
                }
            ]
        },
        {
            "title": "4. 🏢 Halaman Master Kantor (/kantor)",
            "meta": "File: MasterKantor.jsx | API: GET/POST/PUT/DELETE /api/kantor | Controller: KantorController.js | Model: KantorModel.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("master_kantor", "find, findOne, insertOne, updateOne, deleteOne", "Daftar kantor: nopend, nama, kota, tipe (KCU/KCP/AGP/SPP)")
            ]
        },
        {
            "title": "5. 🚛 Halaman Fleet / Kendaraan (/fleet)",
            "meta": "File: MasterKendaraan.jsx | API: GET/POST/PUT/DELETE /api/kendaraan | Controller: KendaraanController.js | Model: KendaraanModel.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("master_kendaraan", "find, findOne, insertOne, updateOne, deleteOne", "List armada, detail kendaraan"),
                ("master_route_nopen", "find({ route_id })", "Rute yang ditugaskan per kendaraan"),
                ("detail_route", "find({ route_id })", "Waypoints rute per kendaraan"),
                ("master_kantor", "find({ nopend })", "Nama kantor stop rute"),
                ("transaksi", "find({ vehicle_nopol })", "Daftar transaksi per armada"),
                ("route_journeys", "findOne({ vehicle_nopol, journey_date })", "Utilisasi kapasitas real-time")
            ]
        },
        {
            "title": "6. 🛤️ Halaman Routes (/routes)",
            "meta": "File: MasterRoute.jsx | API: GET/POST/PUT/DELETE /api/route | Controllers: RouteController.js, DetailRouteController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("master_route_nopen", "find, insertOne, updateOne, deleteOne", "List rute operasional"),
                ("detail_route", "find({ route_id }), insertOne, updateOne", "Waypoints / segmen per rute (seq, asal_nopen, tujuan_nopen)")
            ]
        },
        {
            "title": "7. 📅 Halaman Schedule Templates (/schedule-templates)",
            "meta": "File: TemplateJadwal.jsx | Model: TemplateModel.js → koleksi template_jadwal_transportasi",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("template_jadwal_transportasi", "find, insertOne, updateOne, deleteOne", "Template jadwal default per rute kendaraan")
            ]
        },
        {
            "title": "8. 🗓️ Halaman Transport Schedule (/jadwal-transportasi)",
            "meta": "File: JadwalTransportasi.jsx | Model: JadwalModel.js → koleksi jadwal_transportasi",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("jadwal_transportasi", "find, insertOne, updateOne, deleteOne", "Jadwal transportasi per tanggal operasional"),
                ("master_kendaraan", "findOne", "Data kendaraan terjadwal"),
                ("master_route_nopen", "findOne", "Rute yang dijadwalkan")
            ]
        },
        {
            "title": "9. 🚗 Halaman Milk Run Telemetry / Route Journey (/route-journey)",
            "meta": "File: RouteJourney.jsx | API: GET /api/route-journeys/active | Controller: RouteJourneyController.js | Model: RouteJourneyModel.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("route_journeys", "find/findOne/updateOne", "Status perjalanan aktif, cargo list, current_stop_seq"),
                ("master_kendaraan", "findOne({ nopol })", "Info armada & kapasitas"),
                ("detail_route", "find({ route_id, status: 'AKTIF' })", "Waypoints rute aktif"),
                ("master_kantor", "find({ nopend })", "Nama lokasi tiap stop"),
                ("transaksi", "find({ connote_code })", "Paket yang ada di cargo")
            ]
        },
        {
            "title": "10. 📈 Halaman Estimasi Milk Run (/estimasi-milk-run)",
            "meta": "File: EstimasiMilkRun.jsx | API: GET /api/estimasi-milk-run/... | Controller: EstimasiController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("master_kendaraan", "find({})", "Daftar kendaraan untuk estimasi"),
                ("route_journeys", "findOne({ vehicle_nopol, journey_date })", "Journey aktif per kendaraan"),
                ("detail_route", "find({ route_id })", "Stop-stop rute estimasi"),
                ("master_kantor", "find({ nopend })", "Nama kantor stop"),
                ("transaksi", "find({ vehicle_nopol / cargo })", "Paket dalam hitungan estimasi berat")
            ]
        },
        {
            "title": "11. 🚪 Halaman Gate Monitoring (/gate-monitoring)",
            "meta": "File: GateMonitoring.jsx | API: GET /api/transaksi?... | Controller: TransactionController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("transaksi", "find({ state: 'ENTRY' / 'IN_TRANSIT' })", "Paket yang masuk / keluar gate"),
                ("master_kantor", "find({ nopend })", "Nama kantor asal/tujuan paket"),
                ("master_kendaraan", "findOne({ nopol })", "Info armada pengangkut")
            ]
        },
        {
            "title": "12. 📦 Halaman Transaksi (/transaksi)",
            "meta": "File: Transaksi.jsx | API: GET /api/transaksi | Controllers: TransactionController.js, ManifestController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("transaksi", "find({ filters }), aggregate", "List semua transaksi kiriman dengan filter"),
                ("manifests", "find / findOne", "Data manifest pengiriman"),
                ("master_kantor", "find", "Lookup nama kantor"),
                ("route_journeys", "find", "Journey terkait manifest")
            ]
        },
        {
            "title": "13. 🛍️ Halaman Produk (/produk)",
            "meta": "File: MasterProduk.jsx | Model: ProdukModel.js → koleksi master_produk",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("master_produk", "find, insertOne, updateOne, deleteOne", "Daftar layanan / produk pos (Reguler, Kilat, Express)")
            ]
        },
        {
            "title": "14. 👤 Halaman Profil & User Management (/profil)",
            "meta": "File: Profile.jsx | API: GET/POST/PUT/DELETE /api/users | Controller: UserController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Data Yang Ditampilkan"),
                ("users", "find, insertOne, updateOne, deleteOne", "Manajemen akun: username, nama, role, NIP, cabang")
            ]
        },
        {
            "title": "15. ⚙️ Database Viewer / Compass (/settings atau /compass)",
            "meta": "File: Compass.jsx, Settings.jsx | Controller: CompassController.js",
            "table": [
                ("Koleksi MongoDB", "Operasi", "Keterangan"),
                ("Semua koleksi", "listCollections, find, insertOne, updateOne, deleteOne", "UI admin untuk browse & edit dokumen langsung")
            ]
        }
    ]
    
    def render_table_op(data_tuples):
        t = doc.add_table(rows=len(data_tuples), cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        w = [Inches(2.0), Inches(2.3), Inches(2.5)]
        for r_idx, r in enumerate(t.rows):
            for c_idx, cell in enumerate(r.cells):
                cell.width = w[c_idx]
                cell.text = data_tuples[r_idx][c_idx]
                set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                if r_idx == 0:
                    set_cell_background(cell, "003366")
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9.0)
                else:
                    set_cell_background(cell, "F9FAFB" if r_idx % 2 == 1 else "FFFFFF")
                    for run in p.runs:
                        run.font.size = Pt(8.5)
                        if c_idx == 0:
                            run.font.bold = True

    for p_info in pages_data:
        h = doc.add_heading(level=2)
        rh = h.add_run(p_info["title"])
        rh.font.size = Pt(11)
        rh.font.bold = True
        rh.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        
        p_meta = doc.add_paragraph()
        r_meta = p_meta.add_run(p_info["meta"])
        r_meta.font.size = Pt(9.0)
        r_meta.font.italic = True
        
        if "table" in p_info:
            render_table_op(p_info["table"])
            
        if "subsections" in p_info:
            for sub in p_info["subsections"]:
                p_sub = doc.add_paragraph()
                r_sub = p_sub.add_run(sub["subtitle"])
                r_sub.font.size = Pt(9.5)
                r_sub.font.bold = True
                render_table_op(sub["table"])
                
        if "note" in p_info:
            p_note = doc.add_paragraph()
            r_note = p_note.add_run("Note: " + p_info["note"])
            r_note.font.size = Pt(8.5)
            r_note.font.italic = True

    doc.add_paragraph() # space

    # Section 3: Detailed Field Schema for 12 Collections
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("📋 3. Spesifikasi Rinci Skema Field Per Koleksi MongoDB (12 Koleksi)")
    r3.font.size = Pt(13)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    def add_collection_spec(title_text, fields_data):
        h = doc.add_heading(level=2)
        rh = h.add_run(title_text)
        rh.font.size = Pt(11)
        rh.font.bold = True
        rh.font.color.rgb = RGBColor(0xCC, 0x33, 0x00) # POS Red
        
        headers = ["Nama Field / Key", "Tipe Data", "Contoh Nilai", "Deskripsi & Fungsi Field", "Keterhubungan Ke Tampilan Web (UI & Fitur)"]
        t = doc.add_table(rows=len(fields_data) + 1, cols=5)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = [Inches(1.5), Inches(0.8), Inches(1.1), Inches(1.7), Inches(1.7)]
        
        hdr_cells = t.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].width = widths[idx]
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "003366")
            set_cell_margins(hdr_cells[idx], top=60, bottom=60, left=60, right=60)
            p = hdr_cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(8.5)
                
        for row_idx, f_item in enumerate(fields_data):
            row_cells = t.rows[row_idx + 1].cells
            bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
            for col_idx in range(5):
                cell = row_cells[col_idx]
                cell.width = widths[col_idx]
                cell.text = f_item[col_idx]
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(8.0)
                    if col_idx == 0:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0x00, 0x44, 0x88)
        
        doc.add_paragraph()

    # 1. Transaksi
    transaksi_fields = [
        ("_id", "ObjectId", 'ObjectId("66b9d...")', "Identifier unik dokumen MongoDB.", "Kunci internal dokumen di MongoDB Viewer (/compass)."),
        ("connote_code", "String", '"P20260724000001"', "Nomor resi / connote utama pengiriman.", "Input pencarian di /checker, Kolom Resi tabel /transaksi, ID kargo di modal /route-journey & /gate-monitoring."),
        ("connote.connote_code", "String", '"P20260724000001"', "Duplicate alias nomor resi (sub-dokumen).", "Matching resi pada query legacy findByConnoteCode() di TransactionModel.js."),
        ("connote.connote_booking_code", "String", '"BK-P20260724000001"', "Kode booking transaksi dari aplikasi asal (PosAja/iPOS).", "Tampil sebagai Booking Code di modal detail resi /checker & /transaksi."),
        ("connote.connote_service", "String", '"Pos Reguler", "Pos Express"', "Nama layanan pengiriman pos.", "Badge Jenis Layanan (Kuning/Biru) di /checker & filter dropdown layanan di /transaksi."),
        ("connote.connote_amount", "Number", '35000', "Biaya / ongkos kirim paket (IDR).", "Kolom Biaya Kirim pada tabel transaksi /transaksi & rincian tarif di /checker."),
        ("connote.actual_weight", "Number", '25.5', "Berat fisik riil paket dalam satuan Kilogram (kg).", "Dipakai menghitung current_load_kg & persen muatan armada di /checker, /route-journey, dan /estimasi-milk-run."),
        ("connote.connote_state", "String", '"ENTRY", "IN_TRANSIT", "DELIVERED"', "Status siklus keberadaan fisik paket.", "Badge Warna Status di /checker (Biru: ENTRY, Kuning: IN_TRANSIT, Hijau: DELIVERED) & Ringkasan Stats di / (Dashboard)."),
        ("connote.connote_sender_name", "String", '"PT Pos Logistics Store"', "Nama lengkap pihak pengirim paket.", "Kartu Detail Pengirim di /checker & kolom Pengirim di tabel /transaksi."),
        ("connote.connote_receiver_name", "String", '"SPP Bandung Hub"', "Nama lengkap pihak penerima paket.", "Kartu Detail Penerima di /checker & kolom Penerima di tabel /transaksi."),
        ("connote.connote_receiver_address", "String", '"Jl. Soekarno Hatta No. 564, Bandung"', "Alamat lengkap tujuan pengantaran.", "Tampilan Alamat Penerima pada modal detail resi di /checker."),
        ("connote.created_at", "String/Date", '"24/07/2026 08:00"', "Tanggal dan waktu transaksi dibuat di loket pos.", "Info Tanggal Dibuat di header hasil pencarian /checker & tabel /transaksi."),
        ("location_data_created.location_name", "String", '"Kantor 40511"', "Nama lokasi loket pembuatan transaksi.", "Label lokasi asal pada header tracking /checker."),
        ("location_data_created.custom_field.nopen", "String", '"40511"', "Kode Nopend kantor pos lokasi pembuatan.", "Digunakan backend routing engine untuk memetakan titik asal rute di /checker."),
        ("custom_field.origin_nopen", "String", '"40511"', "Kode Nopend kantor pos pengirim (Origin).", "Titik asal A pada Stepper Rute /checker & pencocokan rute feeder."),
        ("custom_field.destination_nopen", "String", '"40400"', "Kode Nopend kantor pos tujuan (Destination).", "Titik tujuan B pada Stepper Rute /checker & pencocokan rute tujuan.")
    ]
    add_collection_spec("1. Koleksi transaksi (Data Paket / Connote Kiriman)", transaksi_fields)

    # 2. Master Kendaraan
    kendaraan_fields = [
        ("_id", "ObjectId", 'ObjectId("66b9e...")', "Identifier unik dokumen MongoDB.", "Kunci internal dokumen."),
        ("kendaraan_id", "String", '"VH-B9910PCX"', "Kode unik identifier armada kendaraan.", "Relasi ke template_jadwal_transportasi & pilihan dropdown ID di /fleet."),
        ("nopol", "String", '"B 9910 PCX"', "Plat nomor polisi kendaraan (Plat No).", "Kunci pencarian di /checker (Tab List Fleet), Judul Kartu Armada di /fleet, /route-journey, /estimasi-milk-run."),
        ("nama_kendaraan", "String", '"Daihatsu Gran Max Box - Feeder Express"', "Nama deskriptif lengkap kendaraan.", "Subtitle header kendaraan di /fleet & modal detail armada di /route-journey."),
        ("jenis_kendaraan", "String", '"MOBIL BOX INTERCITY (1.5 TON)"', "Klasifikasi jenis & tipe bodi kendaraan.", "Badge Tipe Mobil di /fleet & kartu fleet panel /checker."),
        ("kapasitas_ton", "Number", '1.5', "Kapasitas daya angkut maksimal (Ton).", "Spesifikasi teknik pada kartu detail armada /fleet."),
        ("max_capacity_kg", "Number", '1500', "Kapasitas angkut maksimum dalam kg (Denominator).", "Pembagi rumus % Utilisasi Beban ((load/max)*100) di /checker, /route-journey, /estimasi-milk-run. Mengatur warna Gauge (Hijau <70%, Kuning 70-90%, Oranye 90-100%, Merah >100%)."),
        ("driver", "String", '"Ahmad Supriadi"', "Nama pengemudi / driver penanggung jawab.", "Informasi Driver di kartu /fleet, modal /route-journey, & info armada /checker."),
        ("driver_phone", "String", '"0812-9876-54321"', "Nomor kontak WhatsApp/telepon pengemudi.", "Tombol Hubungi Driver pada modal detail armada di /fleet & /route-journey."),
        ("status", "String", '"AKTIF", "PERBAIKAN"', "Status kelayakan operasional kendaraan.", "Badge Status (Hijau: AKTIF, Merah: PERBAIKAN) di /fleet & filter ketersediaan armada."),
        ("home_base", "String", '"40511 - KCU Cimahi"', "Kantor pangkalan / garasi asal armada.", "Kolom Home Base di /fleet & lokasi pangkalan awal pada pencarian armada /checker."),
        ("rute_utama / assigned_route_id", "String", '"RT-MALAM-B9910-PCX"', "ID rute utama yang ditugaskan ke armada.", "Auto-select rute trayek saat armada dipilih di /route-journey, /jadwal-transportasi, & /estimasi-milk-run.")
    ]
    add_collection_spec("2. Koleksi master_kendaraan (Data Armada Logistik)", kendaraan_fields)

    # 3. Master Kantor
    kantor_fields = [
        ("_id", "ObjectId", 'ObjectId("66b9f...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("nopend", "String", '"40511", "40400"', "Nomor Pendirian / Kode Unik Kantor Pos (5 digit).", "Primary Key Lookup: Digunakan seluruh halaman (/checker, /kantor, /routes, /route-journey) untuk memetakan kode ke nama kantor."),
        ("nama_nopend", "String", '"KCU Cimahi", "SPP Bandung"', "Nama resmi kantor pos.", "Label Nama Lokasi pada Stepper Rute di /checker, header stop /route-journey, & kolom Nama Kantor di /kantor."),
        ("nopen_kc_kcu", "String", '"40500"', "Kode Nopend KCU pembina / kantor utama.", "Hierarki pengelompokan kantor cabang di bawah KCU pada halaman /kantor."),
        ("kdregional", "String", '"05"', "Kode Regional Wilayah Pos Indonesia.", "Filter Wilayah Regional (Regional 5 Jabar) pada tabel /kantor."),
        ("tipe", "String", '"KCU", "KCP", "SPP", "AGEN"', "Jenis/klasifikasi peran kantor pos.", "Badge Tipe Kantor di /kantor & Ikon titik stop pada Stepper Rute /checker (Ikon Hub SPP vs Kantor Cabang)."),
        ("status", "String", '"AKTIF"', "Status operasional kantor pos.", "Filter kantor aktif di /kantor.")
    ]
    add_collection_spec("3. Koleksi master_kantor (Data Kantor Pos / Node Nopend)", kantor_fields)

    # 4. Master Route Nopen
    route_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba0...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("route_id", "String", '"RT-MALAM-B9910-PCX"', "Kode ID unik rute trayek.", "Kunci relasi dengan detail_route, dropdown rute di /routes, /fleet, /jadwal-transportasi, /route-journey."),
        ("nama_route", "String", '"Rute Malam Feeder Cimahi -> SPP Bandung"', "Nama lengkap rute trayek.", "Judul Utama Rute pada tabel /routes, header Stepper /checker & /route-journey."),
        ("nopen_asal", "String", '"40511"', "Kode Nopend kantor titik keberangkatan awal (Origin).", "Teks Origin pada kartu /routes & filter rute berdasarkan titik awal di /checker."),
        ("nama_asal", "String", '"KCU Cimahi (40511)"', "Nama kantor asal keberangkatan.", "Label lokasi asal rute di /routes & /checker."),
        ("nopen_tujuan", "String", '"40400"', "Kode Nopend kantor titik tujuan akhir (Destination).", "Teks Destination pada kartu /routes & filter rute tujuan di /checker."),
        ("nama_tujuan", "String", '"SPP Bandung (40400)"', "Nama kantor tujuan akhir.", "Label lokasi tujuan rute di /routes & /checker."),
        ("kodeMile", "String", '"FIRST_MILE", "MIDDLE_MILE"', "Kategori tahapan distribusi logistik.", "Badge Warna Mile Category (First Mile = Hijau, Middle Mile = Biru, Last Mile = Ungu) di /routes & /checker."),
        ("deskripsi_produk", "String", '"Pos Reguler & Express Pickup Malam"', "Keterangan jenis muatan produk di rute.", "Kolom Deskripsi di tabel /routes."),
        ("prioritas", "Number", '1, 2', "Tingkat prioritas pemilihan rute.", "Dipakai algoritma backend checkRouting() untuk memilih rute paling utama saat pencarian resi di /checker."),
        ("status_route", "String", '"LENGKAP"', "Status kelengkapan struktur rute.", "Badge status kelengkapan data rute di /routes."),
        ("aktif", "String", '"Y", "N"', "Flag keaktifan rute operasional.", "Toggle status Aktif/Nonaktif di modal /routes.")
    ]
    add_collection_spec("4. Koleksi master_route_nopen (Header Rute Statis Operasional)", route_fields)

    # 5. Detail Route
    detail_route_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba1...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("detail_route_id", "String", '"DR-B9910-01"', "ID unik per titik segmen rute.", "Referensi internal editing waypoint di /routes."),
        ("route_id", "String", '"RT-MALAM-B9910-PCX"', "ID rute induk pemegang waypoint.", "Foreign Key ke master_route_nopen. Backend menarik seluruh waypoint rute ini untuk menyusun Stepper di /checker & /route-journey."),
        ("seq", "Number", '1, 2, 3, 4, 5', "Nomor urut persinggahan (Sequence Stop).", "Menentukan urutan nomor lingkaran (Stop 1, Stop 2, dst) pada Stepper Rute di /checker & /route-journey."),
        ("asal_nopen", "String", '"40511"', "Kode Nopend kantor awal segmen ini.", "Label asal segmen waypoint di Stepper /checker."),
        ("asal_nama", "String", '"KCU Cimahi"', "Nama kantor asal segmen ini.", "Teks nama kantor pada node Stepper /checker."),
        ("tujuan_nopen", "String", '"40521"', "Kode Nopend kantor tujuan segmen ini.", "Label tujuan segmen waypoint di Stepper /checker."),
        ("tujuan_nama", "String", '"KCP Cimahi Selatan"', "Nama kantor tujuan segmen ini.", "Teks nama kantor persinggahan pada node Stepper di /checker & /route-journey."),
        ("estimasi_menit", "Number", '12, 25', "Waktu tempuh segmen dalam menit.", "Menampilkan durasi antar-stop (e.g. ⏱️ 12 min) pada garis Stepper /checker & total ETA di /estimasi-milk-run."),
        ("jarak_km", "Number", '5.2, 8.5', "Jarak fisik segmen dalam Kilometer (km).", "Menampilkan jarak antar-stop (e.g. 📏 5.2 km) pada garis Stepper /checker & akumulasi total KM rute di /routes."),
        ("status", "String", '"AKTIF"', "Status keaktifan segmen waypoint.", "Filter waypoint aktif di /routes.")
    ]
    add_collection_spec("5. Koleksi detail_route (Waypoint / Stop Segmen Rute)", detail_route_fields)

    # 6. Route Journeys
    journeys_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba2...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("journey_id", "String", '"JRN-20260724-B9910PCX-001"', "Kode unik transaksi perjalanan operasional armada.", "Header utama halaman /route-journey, parameter URL API, & log audit operasional."),
        ("vehicle_nopol", "String", '"B 9910 PCX"', "Plat nomor armada yang menjalankan journey.", "Matching journey dengan armada di /checker & /route-journey."),
        ("route_id", "String", '"RT-MALAM-B9910-PCX"', "ID rute yang sedang ditempuh.", "Menentukan daftar waypoints yang ditampilkan pada Stepper Live /route-journey."),
        ("status", "String", '"IN_PROGRESS", "COMPLETED"', "Status eksekusi perjalanan armada.", "Badge Warna Status di /route-journey & /checker (Kuning: IN_PROGRESS, Hijau: COMPLETED) & statistik aktif di / (Dashboard)."),
        ("current_stop_seq", "Number", '1, 2, 3', "Nomor urut stop tempat armada saat ini berada/terakhir singgah.", "Penggerak Stepper Live: Mengatur status node di /checker & /route-journey (seq < current_stop_seq → COMPLETED/Centang Hijau, seq === current_stop_seq → CURRENT/Mobil Berkedip, seq > current_stop_seq → UPCOMING/Abu-abu)."),
        ("maximum_capacity_kg", "Number", '1500', "Batas maksimum muatan mobil (kg).", "Pembagi persen muatan pada Gauge Meter /route-journey."),
        ("current_load_kg", "Number", '619.0', "Total berat muatan paket aktif yang ada di atas mobil (kg).", "Meter Utilisasi Muatan Real-time di /route-journey & /checker."),
        ("shift", "String", '"MALAM", "SIANG"', "Shift jadwal kerja operasional.", "Badge Shift di header /route-journey."),
        ("tanggal_operasional", "Date/String", '"2026-07-24"', "Tanggal pelaksanaan perjalanan.", "Picker Filter Tanggal di /checker & /route-journey."),
        ("cargo", "Array", '[{ connote_code: "P...", weight_kg: 25.5 }]', "List paket/barang yang dimuat di dalam armada.", "Tabel Kargo Aktif Dalam Mobil di /route-journey & detail muatan armada saat pencarian nopol di /checker."),
        ("cargo[].connote_code", "String", '"P20260724000001"', "Resi paket di dalam muatan armada.", "Kolom Resi di tabel kargo /route-journey."),
        ("cargo[].weight_kg", "Number", '25.5', "Berat paket individu (kg).", "Kolom Berat Paket di tabel kargo /route-journey."),
        ("cargo[].origin_nopen", "String", '"40511"', "Kode Nopend kantor asal muat paket.", "Kolom Asal Kargo di /route-journey."),
        ("cargo[].destination_nopen", "String", '"40400"', "Kode Nopend kantor tujuan bongkar paket.", "Kolom Tujuan Kargo di /route-journey."),
        ("cargo[].loaded_at_seq", "Number", '1', "Stop sequence tempat paket di-LOAD ke mobil.", "Log urutan muat kargo di /route-journey."),
        ("cargo[].unloaded_at_seq", "Number", '6', "Stop sequence tempat paket akan di-UNLOAD dari mobil.", "Log urutan bongkar kargo di /route-journey.")
    ]
    add_collection_spec("6. Koleksi route_journeys (Perjalanan Harian Armada & Kargo Telemetri)", journeys_fields)

    # 7. Jadwal Transportasi
    jadwal_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba3...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("jadwal_id", "String", '"SCH-20260724-001"', "Kode unik record jadwal harian.", "Referensi internal tabel /jadwal-transportasi."),
        ("route_id", "String", '"RT-MALAM-B9910-PCX"', "ID rute yang dijadwalkan.", "Kolom Rute di /jadwal-transportasi & matching jadwal di /checker."),
        ("nomor_polisi / nopol", "String", '"B 9910 PCX"', "Plat nomor armada yang ditugaskan.", "Kolom Armada / Plat No di /jadwal-transportasi & /checker."),
        ("tanggal", "String/Date", '"2026-07-24"', "Tanggal berlakunya jadwal operasional.", "Filter Tanggal di halaman /jadwal-transportasi."),
        ("jam_berangkat", "String", '"21:00"', "Rencana jam keberangkatan dari origin (ETD).", "Kolom Jam Berangkat di /jadwal-transportasi & info ETD di /checker."),
        ("jam_tiba", "String", '"03:30"', "Rencana jam kedatangan di tujuan (ETA).", "Kolom Jam Tiba di /jadwal-transportasi & info ETA di /checker."),
        ("status", "String", '"TERJADWAL", "BERJALAN"', "Status pelaksanaan jadwal.", "Badge warna status di tabel /jadwal-transportasi."),
        ("sumber_generate", "String", '"AUTO_GENERATE_TEMPLATE"', "Catatan asal mula jadwal.", "Kolom Keterangan / Sumber di /jadwal-transportasi.")
    ]
    add_collection_spec("7. Koleksi jadwal_transportasi (Jadwal Operasional Harian)", jadwal_fields)

    # 8. Template Jadwal
    tmpl_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba4...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("template_id", "String", '"TMPL-RT-MALAM-B9910"', "ID unik template jadwal.", "Referensi tabel di /schedule-templates."),
        ("route_id", "String", '"RT-MALAM-B9910-PCX"', "ID rute rujukan template.", "Dropdown Rute di /schedule-templates & fallback lookup jadwal di backend /checker."),
        ("kendaraan_id", "String", '"VH-B9910PCX"', "ID armada default untuk rute ini.", "Kolom Armada Default di /schedule-templates."),
        ("jam_berangkat", "String", '"21:00"', "Jam berangkat standar harian.", "Kolom Jam Berangkat Default di /schedule-templates."),
        ("jam_tiba", "String", '"03:30"', "Jam tiba standar harian.", "Kolom Jam Tiba Default di /schedule-templates."),
        ("hari_operasional", "Array/String", '["SENIN", "SELASA"]', "Hari-hari berlakunya jadwal mingguan.", "Checklist Hari Operasional pada modal /schedule-templates."),
        ("status", "String", '"AKTIF"', "Status keaktifan template jadwal.", "Toggle status di /schedule-templates.")
    ]
    add_collection_spec("8. Koleksi template_jadwal_transportasi (Template Jadwal Default)", tmpl_fields)

    # 9. Manifests
    manifest_fields = [
        ("master_manifest_code", "String", '"MF-20260724-B9910PCX"', "Kode barcode / nomor unik manifest kantong pos.", "Header Manifest pada tab Kantong /transaksi & cetak lembar serah terima manifest."),
        ("asal_nopen", "String", '"40511"', "Kode Nopend kantor yang menyegel manifest.", "Tampilan Kantor Asal Manifest di /transaksi."),
        ("tujuan_nopen", "String", '"40400"', "Kode Nopend kantor tujuan kantong.", "Tampilan Kantor Tujuan Manifest di /transaksi."),
        ("total_connote", "Number", '10', "Total jumlah paket di dalam kantong.", "Kolom Total Paket pada tabel manifest /transaksi."),
        ("total_weight_kg", "Number", '619.0', "Total berat akumulasi kantong (kg).", "Kolom Total Berat (kg) pada tabel manifest /transaksi."),
        ("status", "String", '"SEALED", "IN_TRANSIT"', "Status fisik kantong pos.", "Badge status manifest di /transaksi."),
        ("created_by", "String", '"SUPER_ADMIN"', "Petugas pembuat manifest kantong.", "Kolom Petugas di /transaksi."),
        ("manifest_detail.connote_code", "String", '"P20260724000001"', "Resi paket individual di dalam kantong.", "Item list resi saat baris manifest di-expand pada /transaksi.")
    ]
    add_collection_spec("9. Koleksi manifest_master & manifest_detail (Manifest Kantong Pos)", manifest_fields)

    # 10. Tracking Events
    events_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba5...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("event_id", "String", '"P20260724000001_LOADED_..."', "Key unik penjamin idempotensi (cegah duplikasi saat re-import CSV).", "Deduplikasi backend & log audit."),
        ("connote_code", "String", '"P20260724000001"', "Resi paket yang discan.", "Foreign Key query timeline riwayat tracking di /checker."),
        ("event_type", "String", '"ENTRY", "LOADED", "DELIVERED"', "Kode jenis kejadian fisik logistik.", "Penyusun Timeline UI: Menentukan Ikon Status & Judul Langkah pada Timeline Riwayat Tracking /checker (ENTRY = 📦 Terima di Loket, LOADED = 🚚 Dimuat ke Armada, ARRIVED = 🏢 Tiba di Node, DELIVERED = ✅ Selesai)."),
        ("event_datetime", "Date/String", '2026-07-24T10:15:00Z', "Waktu persis scan dilakukan.", "Tampilan Jam & Tanggal pada setiap baris Timeline /checker."),
        ("office_code / location_name", "String", '"40511", "KCU Cimahi"', "Lokasi kantor pos saat scan.", "Label Lokasi pada baris Timeline /checker."),
        ("vehicle_code", "String", '"B 9910 PCX"', "Plat nomor armada pengangkut saat scan.", "Teks Keterangan \"Diangkut oleh B 9910 PCX\" pada event LOADED di /checker."),
        ("stop_sequence", "Number", '1, 2', "Urutan stop tempat scan terjadi.", "Menyorot node stop yang bersesuaian di Stepper /checker."),
        ("import_batch_id", "String", '"BATCH-20260724-153045-8A2F"', "ID batch impor CSV harian.", "Tagging audit data untuk fitur komited/rollback impor CSV di /checker.")
    ]
    add_collection_spec("10. Koleksi tracking_events (Event Log Tracking Scan Paket)", events_fields)

    # 11. Users
    users_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba6...")', "Identifier unik dokumen user.", "Internal MongoDB."),
        ("username", "String", '"admin", "sari", "dispatcher"', "Nama akun unik untuk login.", "Input Username di form /login & daftar pengguna di /profil."),
        ("password_hash", "String", '"$2b$10$3pZ51LdGg9g..."', "Password terenkripsi Bcrypt.", "Otentikasi keamanan saat tombol Login ditekan di /login."),
        ("name", "String", '"Sari Rahayu"', "Nama lengkap pengguna.", "Tampilan Nama Pengguna di Pojok Kanan Atas App Bar & Kartu Profil /profil."),
        ("role", "String", '"SUPER_ADMIN", "DISPATCHER"', "Peran / Hak Akses Pengguna (RBAC).", "Proteksi Fitur Web: Membatasi menu sidebar & proteksi route halaman React (Super Admin = full akses, Driver = route journey saja)."),
        ("NIP / branch", "String", '"99283741", "40511 - KCU Cimahi"', "Nomor Induk Pegawai & Cabang Tugas.", "Informasi NIP & Cabang pada halaman Profil /profil.")
    ]
    add_collection_spec("11. Koleksi users (Data Akun Pengguna & Role RBAC)", users_fields)

    # 12. Master Produk
    produk_fields = [
        ("_id", "ObjectId", 'ObjectId("66ba7...")', "Identifier unik dokumen.", "Internal MongoDB."),
        ("serviceId", "String", '"POS-REGULER", "POS-NEXTDAY"', "Kode unik produk layanan pos.", "Dropdown Layanan saat input transaksi di /transaksi & tabel katalog /produk."),
        ("kodeMile", "String", '"FIRST_MILE", "MIDDLE_MILE"', "Segmen jangkauan distribusi pos.", "Badge Mile pada katalog layanan /produk."),
        ("deskripsi", "String", '"Pos Reguler Pengiriman Standar Intercity"', "Penjelasan rinci SLA layanan.", "Kolom Deskripsi di tabel /produk."),
        ("segmenProduk", "String", '"REGULER", "EXPRESS"', "Kategori kelas kecepatan kirim.", "Filter Segmen Produk pada /produk."),
        ("pasar", "String", '"DOMESTIK"', "Lingkup pasar layanan (Domestik/Intl).", "Tag Pasar pada /produk."),
        ("status", "String", '"AKTIF"', "Status keaktifan produk.", "Toggle status layanan di /produk.")
    ]
    add_collection_spec("12. Koleksi master_produk (Katalog Layanan Produk Pos)", produk_fields)

    # Save output file
    output_path = r"c:\Users\Asus\Documents\POSIND\IPOS5\ipos5\db_integration_docs.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX file at: {output_path}")

if __name__ == '__main__':
    create_document()
