import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUTPUT_DIR = r"c:\Users\Asus\Documents\POSIND\IPOS5\ipos5\dokumentasi_halaman_ipos5"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_page_documentation(filename_base, title, metadata, sections):
    md_path = os.path.join(OUTPUT_DIR, f"{filename_base}.md")
    docx_path = os.path.join(OUTPUT_DIR, f"{filename_base}.docx")

    # 1. WRITE MARKDOWN (.md) FILE
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# DOKUMENTASI TEKNIS & OPERASIONAL HALAMAN\n")
        f.write(f"## {title.upper()}\n\n")
        f.write("### 📌 METADATA & INTEGRASI SISTEM\n")
        f.write("| Parameter | Nilai / Spesifikasi |\n")
        f.write("| :--- | :--- |\n")
        for k, v in metadata:
            f.write(f"| **{k}** | {v} |\n")
        f.write("\n---\n\n")

        for sec_title, sec_content in sections:
            f.write(f"## {sec_title}\n\n")
            f.write(f"{sec_content}\n\n")

    print(f"[OK] Generated MD: {filename_base}.md")

    # 2. WRITE WORD DOCUMENT (.docx) FILE
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # Title Banner
    tp = doc.add_paragraph()
    tr = tp.add_run(f"DOKUMENTASI TEKNIS HALAMAN\n{title.upper()}")
    tr.font.name = 'Arial'
    tr.font.size = Pt(18)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(13, 27, 56) # Navy
    tp.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sr = sub.add_run("IPOS5 Routing & Schedule Management System — PT Pos Indonesia (Persero)")
    sr.font.name = 'Arial'
    sr.font.size = Pt(10)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor(232, 67, 31) # Pos Orange
    sub.paragraph_format.space_after = Pt(14)

    # Metadata Table
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, (k, v) in enumerate(metadata):
        row = table.rows[idx]
        ck, cv = row.cells[0], row.cells[1]
        ck.width = Inches(2.2)
        cv.width = Inches(4.5)
        set_cell_background(ck, "0D1B38")
        set_cell_background(cv, "F8FAFC")
        set_cell_margins(ck)
        set_cell_margins(cv)

        pk = ck.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.name = 'Arial'
        rk.font.size = Pt(9.5)
        rk.font.bold = True
        rk.font.color.rgb = RGBColor(255, 255, 255)

        pv = cv.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.name = 'Arial'
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Add Sections to DOCX
    for sec_title, sec_content in sections:
        hp = doc.add_paragraph()
        hr = hp.add_run(sec_title)
        hr.font.name = 'Arial'
        hr.font.size = Pt(13)
        hr.font.bold = True
        hr.font.color.rgb = RGBColor(2, 132, 199) # Sky Blue
        hp.paragraph_format.space_before = Pt(14)
        hp.paragraph_format.space_after = Pt(6)

        for line in sec_content.split('\n'):
            line_s = line.strip()
            if not line_s:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15

            if line_s.startswith('- ') or line_s.startswith('* '):
                p.paragraph_format.left_indent = Inches(0.25)
                rb = p.add_run("• ")
                rb.font.bold = True
                rb.font.color.rgb = RGBColor(232, 67, 31)
                
                rt = p.add_run(line_s[2:])
                rt.font.name = 'Arial'
                rt.font.size = Pt(9.5)
                rt.font.color.rgb = RGBColor(51, 65, 85)
            elif line_s.startswith('### '):
                rs = p.add_run(line_s[4:])
                rs.font.name = 'Arial'
                rs.font.size = Pt(10.5)
                rs.font.bold = True
                rs.font.color.rgb = RGBColor(13, 27, 56)
                p.paragraph_format.space_before = Pt(8)
            else:
                rt = p.add_run(line_s)
                rt.font.name = 'Arial'
                rt.font.size = Pt(9.5)
                rt.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(docx_path)
    print(f"[OK] Generated DOCX: {filename_base}.docx")


# ==============================================================================
# DATA DEFINITIONS FOR ALL 16 PAGES
# ==============================================================================

PAGES_DATA = [

    # PAGE 00: INDEX
    {
        "filename": "00_INDEX_DOKUMENTASI_ALL_PAGES",
        "title": "Master Index Dokumentasi Seluruh Halaman IPOS5",
        "metadata": [
            ("Aplikasi Induk", "IPOS5 Routing & Schedule Management System"),
            ("Organisasi", "PT Pos Indonesia (Persero)"),
            ("Total Modul / Halaman", "15 Modul Utama (16 File Dokumentasi)"),
            ("Arsitektur", "React Vite + Express.js OOP MVC + MongoDB Multi-Server"),
            ("Versi Sistem", "Redesign Enterprise v2.5.0")
        ],
        "sections": [
            ("1. Pendahuluan & Struktur Dokumentasi", 
             "Dokumentasi ini mencakup penjelasan teknis dan operasional secara detail untuk seluruh 15 modul halaman pada aplikasi IPOS5 PT Pos Indonesia. Setiap halaman didokumentasikan dalam 2 format file independen: format Markdown (.md) dan Microsoft Word (.docx).\n\n"
             "Sistem IPOS5 dirancang untuk mengelola pendistribusian kiriman pos, manajemen armada kendaraan feeder/intercity, penjadwalan rantai pasok logistik, telemetri Milk Run multi-stop, dan pemantauan gate transit gerbang gerak secara real-time."),
            
            ("2. Daftar Pemetaan Modul Halaman & Route URL",
             "Berikut adalah daftar lengkap 15 modul halaman aplikasi IPOS5 beserta URL route dan file komponen React pendukung:\n\n"
             "- 1. Dashboard Operasional (`/`) -> File: `Dashboard.jsx`\n"
             "- 2. Package Tracking & Checker (`/checker`) -> File: `Checker.jsx`\n"
             "- 3. Master Kantor Pos (`/kantor`) -> File: `MasterKantor.jsx`\n"
             "- 4. Master Produk & Layanan (`/produk`) -> File: `MasterProduk.jsx`\n"
             "- 5. Master Kendaraan Armada (`/kendaraan`) -> File: `MasterKendaraan.jsx`\n"
             "- 6. Master Route Logistik (`/route`) -> File: `MasterRoute.jsx`\n"
             "- 7. Template Jadwal Transportasi (`/template`) -> File: `TemplateJadwal.jsx`\n"
             "- 8. Transport Schedule Harian (`/jadwal`) -> File: `JadwalTransportasi.jsx`\n"
             "- 9. Milk Run Telemetry & Execution (`/route-journey`) -> File: `RouteJourney.jsx`\n"
             "- 10. Estimasi Milk Run (`/estimasi`) -> File: `EstimasiMilkRun.jsx`\n"
             "- 11. Gate Monitoring & Transit Hub (`/transit-monitoring`) -> File: `GateMonitoring.jsx`\n"
             "- 12. Analitik & Laporan Operasional (`/analytics`) -> File: `AnalyticsReport.jsx`\n"
             "- 13. Database Viewer Transaksi (`/transaksi`) -> File: `Transaksi.jsx`\n"
             "- 14. System Settings & Compass Connection (`/settings`) -> File: `SettingsPage.jsx`\n"
             "- 15. Profile User & Account (`/profile`) -> File: `Profile.jsx`"),

            ("3. Standar Arsitektur Integrasi Backend & Database",
             "Setiap modul terkoneksi dengan Backend Node.js Express.js bertipe OOP Controller (`TransactionController`, `KendaraanController`, `RouteJourneyController`, `KantorController`, dll.) dan menyimpan/mengambil data dari MongoDB koleksi resmi (`transaksi`, `route_journeys`, `master_kendaraan`, `master_kantor`, `detail_route`, `master_route_nopen`, `jadwal_transportasi`).")
        ]
    },

    # PAGE 01: DASHBOARD
    {
        "filename": "01_Dashboard_Operasional",
        "title": "Halaman 1 — Dashboard Operasional Logistik",
        "metadata": [
            ("Nama Modul", "Dashboard Operasional Logistik IPOS5"),
            ("Kategori Menu", "OPERATIONS"),
            ("Route URL", "/"),
            ("File Komponen React", "src/pages/Dashboard.jsx"),
            ("Backend API Endpoint", "GET /api/transaksi/stats, GET /api/kendaraan/stats, GET /api/routes/stats"),
            ("Koleksi MongoDB", "transaksi, master_kendaraan, route_journeys, master_route_nopen"),
            ("Hak Akses Role", "Semua Role (Operator & Super Admin)")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Halaman Dashboard Operasional merupakan pusat komando (command center) visual bagi manajemen logistik PT Pos Indonesia. Halaman ini menyajikan rangkuman metrik statistik real-time mengenai status pengiriman resi paket, ketersediaan armada mobil, utilisasi rute feeder/intercity, dan status koneksi database MongoDB."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Ringkasan Stat Snapshot Cards\n"
             "- Total Paket Terdaftar (ENTRY / IN_TRANSIT / DELIVERED)\n"
             "- Total Tonase Muatan Aktif (Kilogram / Ton)\n"
             "- Total Armada Beroperasi & Status Utilisasi Kapasitas\n"
             "- Total Rute Feeder & Intercity Aktif\n\n"
             "### Visual Widget & Diagram Telemetri\n"
             "- Diagram Pie Status Pengiriman Kiriman Pos (ENTRY, IN_TRANSIT, DELIVERED)\n"
             "- Bar Chart Utilisasi Kapasitas Kendaraan Utama (B 9910 PCX, B 9945 PCY, dll.)\n"
             "- Table Feeder Route Snapshot & Active Journeys Status\n"
             "- Status Live Database Connection Compass Badge (Green/Red indicator)"),

            ("3. Aturan Logika Bisnis & Penanganan Data",
             "- Data agregasi dihitung secara dinamis dari koleksi MongoDB transaksi dan route_journeys.\n"
             "- Menggunakan polling otomatis / trigger refresh untuk memperbarui angka statistik saat terjadi aktivitas bongkar-muat paket di lapangan."),

            ("4. Alur Penggunaan Operator (User Flow)",
             "- 1. Operator masuk ke sistem IPOS5 dan secara otomatis diarahkannya ke halaman Dashboard (`/`).\n"
             "- 2. Operator meninjau ringkasan metrik volume paket dan armada yang sedang berjalan.\n"
             "- 3. Operator dapat mengeklik salah satu kartu statistik untuk berpindah cepat ke halaman rincian (misal klik Paket -> Package Tracking, klik Armada -> Master Kendaraan).")
        ]
    },

    # PAGE 02: PACKAGE TRACKING / CHECKER
    {
        "filename": "02_Package_Tracking_Checker",
        "title": "Halaman 2 — Package Tracking & Checker Logistik",
        "metadata": [
            ("Nama Modul", "Package Tracking & Checker (/checker)"),
            ("Kategori Menu", "OPERATIONS"),
            ("Route URL", "/checker"),
            ("File Komponen React", "src/pages/Checker.jsx"),
            ("Backend API Endpoint", "GET /api/checker/:connoteCode?date=YYYY-MM-DD, GET /api/checker/vehicle/:nopol"),
            ("Koleksi MongoDB", "transaksi, master_kendaraan, route_journeys, detail_route, master_kantor, tracking_events"),
            ("Hak Akses Role", "Semua Role (Operator & Super Admin)")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Halaman Package Tracking & Checker merupakan modul utama audit trail dan pelacakan lintasan pengiriman paket serta armada kendaraan. Modul ini memungkinkan operator melacak status resi connote (14/15 digit) atau nomor plat polisi armada (misal B 9910 PCX) berdasarkan konteks Tanggal Operasional harian."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Toggle & Search Bar Multi-Fungsi\n"
             "- Mode Pencarian: Nomor Resi Connote vs Daftar Fleet Mobil Armada (MongoDB)\n"
             "- Date Picker Context Tanggal Operasional Harian (YYYY-MM-DD)\n"
             "- Quick Test Sample Chips (Resi P260812000001, P20260724000001, Armada B 9910 PCX)\n\n"
             "### Visual Cards & Telemetri Real-Time\n"
             "- Card Warning Tanggal Operasional Belum Tiba (Warna Merah Prominen tanpa muatan palsu)\n"
             "- Card Overview Spesifikasi Kendaraan / Detail Resi Paket\n"
             "- Tracking Timeline Event Log dari Database (ENTRY, LOADED, IN_TRANSIT, DELIVERED)\n"
             "- Multi-stop Route Journey Stepper (Sequential Waypoints dari detail_route)\n"
             "- Vehicle Capacity Gauge (% Utilisasi Beban, Load at Stop kg vs Max Capacity 1500 kg)\n"
             "- Radar GPS Telemetry Live Modal (`LiveGpsMapModal.jsx`)\n"
             "- QA / Dev CSV Import Modal (`CsvImportModal.jsx`)"),

            ("3. Aturan Bisnis & Logika Penanganan Tanggal Masa Depan",
             "⚠️ **ATURAN TANGGAL MASA DEPAN (STRICT FUTURE DATE RULE)**:\n"
             "- Apabila pengguna memilih tanggal operasional di masa depan yang belum tiba (misal hari ini 14 Agustus 2026 dan memilih 17/22/25/28 Agustus 2026), backend TIDAK AKAN LAGI memuat data perjalanan lampau (fallback) atau membuat muatan palsu (3 pcs / 19.7 kg).\n"
             "- Backend mengembalikan `isFutureDate: true` dan `milk_run: null`.\n"
             "- Frontend menampilkan Banner Peringatan Merah ('Tanggal Operasional Belum Tiba') dan secara otomatis MENYEMBUNYIKAN statistik muatan lampau, progress rute *in-progress* palsu, dan gauge utilisasi."),

            ("4. Alur Penggunaan Operator (User Flow)",
             "- 1. Pengguna memasukkan nomor resi atau memilih plat mobil armada.\n"
             "- 2. Pengguna mementukan Tanggal Operasional (misal 12 Agustus 2026).\n"
             "- 3. Sistem memuat lintasan rute, posisi stop aktif, serta daftar muatan barang secara real-time.")
        ]
    },

    # PAGE 03: POST OFFICES / MASTER KANTOR
    {
        "filename": "03_Post_Offices_Master_Kantor",
        "title": "Halaman 3 — Post Offices (Master Data Kantor Pos)",
        "metadata": [
            ("Nama Modul", "Master Kantor Pos (/kantor)"),
            ("Kategori Menu", "OPERATIONS / MASTER DATA"),
            ("Route URL", "/kantor"),
            ("File Komponen React", "src/pages/MasterKantor.jsx"),
            ("Backend API Endpoint", "GET /api/kantor, POST /api/kantor, PUT /api/kantor/:id, DELETE /api/kantor/:id"),
            ("Koleksi MongoDB", "master_kantor"),
            ("Hak Akses Role", "Semua Role (CRUD Terbatas untuk Operator)")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Halaman Master Kantor Pos berfungsi untuk mengelola seluruh direktori titik node jaringan kantor pos di Indonesia, mulai dari Sentral Pengolahan Pos (SPP), Kantor Cabang Utama (KCU), Kantor Cabang (KC), Kantor Cabang Pembantu (KCP), hingga Agen Pos."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Manajemen Data & Filter Node\n"
             "- Tabel Master Kantor Pos dengan Pagination & Search Nopend / Nama Kantor\n"
             "- Badge Klasifikasi Tipe Kantor (SPP = Biru, KCU = Ungu, KC = Hijau, KCP = Oranye, AGEN = Abu-abu)\n"
             "- Filter Dropdown berdasarkan Regional (Reg 1 - Reg 6 Jawa & Luar Jawa)\n"
             "- Modal Tambah & Edit Data Kantor Pos (Nopend 5 digit, Nopen Induk KCU/KC, Alamat, Koordinat Lat/Long)\n"
             "- Fitur Export Data Kantor Pos ke CSV / Excel"),

            ("3. Aturan Logika Bisnis",
             "- Nopend (Nomor Pendirian) bersifat unik (Primary Key Lookup) 5 digit angka.\n"
             "- Nopend digunakan oleh routing engine backend (`TransactionController.checkRouting`) untuk memetakan titik asal (Origin) dan tujuan (Destination) pengiriman resi.")
        ]
    },

    # PAGE 04: PRODUCTS & SERVICES / MASTER PRODUK
    {
        "filename": "04_Products_Services_Master_Produk",
        "title": "Halaman 4 — Products & Services (Master Produk Layanan)",
        "metadata": [
            ("Nama Modul", "Master Produk & Layanan Pos (/produk)"),
            ("Kategori Menu", "MASTER DATA"),
            ("Route URL", "/produk"),
            ("File Komponen React", "src/pages/MasterProduk.jsx"),
            ("Backend API Endpoint", "GET /api/produk, POST /api/produk, PUT /api/produk/:id, DELETE /api/produk/:id"),
            ("Koleksi MongoDB", "master_produk"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul ini mengelola katalog produk layanan pengiriman kurir dan logistik Pos Indonesia (seperti Pos Express, Pos Reguler, Pos Nextday, Pos Jumbo, Kakap, dll.) beserta batasan SLA (Service Level Agreement), batas berat maksimum, serta waktu penutupan loket (Cut-Off Time)."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Katalog Layanan & Konfigurasi SLA\n"
             "- Grid & Table Katalog Layanan Pos dengan Indikator Warna Badge Produk\n"
             "- Pengaturan Estimasi Pengiriman (SLA Hari / Jam H+1, H+2)\n"
             "- Pengaturan Jam Cut-Off Loket Pemrosesan Pengiriman\n"
             "- Form Input Produk Baru & Pengaturan Tarif Dasar Per Kg\n"
             "- Toggle Status Aktif / Non-Aktif Layanan Pos")
        ]
    },

    # PAGE 05: FLEET / MASTER KENDARAAN
    {
        "filename": "05_Fleet_Master_Kendaraan",
        "title": "Halaman 5 — Fleet Management (Master Kendaraan Armada)",
        "metadata": [
            ("Nama Modul", "Master Kendaraan Armada (/kendaraan)"),
            ("Kategori Menu", "MASTER DATA"),
            ("Route URL", "/kendaraan"),
            ("File Komponen React", "src/pages/MasterKendaraan.jsx"),
            ("Backend API Endpoint", "GET /api/kendaraan, POST /api/kendaraan, PUT /api/kendaraan/:id, DELETE /api/kendaraan/:id"),
            ("Koleksi MongoDB", "master_kendaraan, route_journeys"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Halaman Fleet Management bertanggung jawab mengelola seluruh unit kendaraan operasional angkutan pos (Gran Max Box, Isuzu Elf Box, Mitsubishi Canter, Hino Wingbox, Blind Van). Modul ini menyimpan nomor polisi (Plat Nopol), nama driver, nomor HP driver, kapasitas angkut maksimum (kg/ton), kantor pangkalan (Home Base), serta rute utama yang ditugaskan."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Kartu Profil Armada & Live Status\n"
             "- Grid Kartu Kendaraan dengan Indikator Kapasitas Maksimum (misal 1500 kg / 4000 kg)\n"
             "- Status Keberadaan Kendaraan Real-Time di MongoDB (AKTIF, IN_PROGRESS, MAINTENANCE)\n"
             "- Modal Input & Edit Kendaraan Baru (Plat Nopol, Tipe Bodi, Driver, Phone, Home Base, Assigned Route)\n"
             "- Filter Berdasarkan Home Base Kantor Pos (40511 Cimahi, 40000 SPP Bandung, dll.)\n"
             "- Akses Cepat Tombol 'Lacak di Checker' untuk melihat muatan live kendaraan tersebut")
        ]
    },

    # PAGE 06: ROUTES / MASTER ROUTE
    {
        "filename": "06_Routes_Master_Route",
        "title": "Halaman 6 — Routes Management (Master Rute Logistik)",
        "metadata": [
            ("Nama Modul", "Master Route Logistik (/route)"),
            ("Kategori Menu", "MASTER DATA"),
            ("Route URL", "/route"),
            ("File Komponen React", "src/pages/MasterRoute.jsx"),
            ("Backend API Endpoint", "GET /api/routes, POST /api/routes, GET /api/routes/:routeId/segments"),
            ("Koleksi MongoDB", "master_route_nopen, detail_route"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul Master Route mengelola skema rute jaringan trayek antar-kantor pos. Setiap rute memiliki identitas `route_id` (contoh: `RT-MALAM-B9910-PCX`), kantor asal (nopen_asal), kantor tujuan (nopen_tujuan), prioritas pemilihan, kategori mile (FIRST_MILE, MIDDLE_MILE, LAST_MILE), serta sekumpulan segmen perhentian (detail_route waypoints)."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Pengelolaan Header Rute & Waypoints Segmen\n"
             "- Tabel Master Route Nopen dengan Informasi Akumulasi Total Jarak (KM) & Total Est. Menit\n"
             "- Modal Detail Waypoints Segmen (Sequence Stop 1, Stop 2... Stop N)\n"
             "- Editor Segmen Rute (Asal Nopen -> Tujuan Nopen, Jarak KM, Estimasi Menit, Role Stop)\n"
             "- Badge Kategori Mile (First Mile = Hijau, Middle Mile = Biru, Last Mile = Ungu)\n"
             "- Toggle Status Aktif Rute (Y/N)")
        ]
    },

    # PAGE 07: SCHEDULE TEMPLATES
    {
        "filename": "07_Schedule_Templates",
        "title": "Halaman 7 — Schedule Templates (Template Jadwal Transportasi)",
        "metadata": [
            ("Nama Modul", "Template Jadwal Transportasi (/template)"),
            ("Kategori Menu", "LOGISTICS"),
            ("Route URL", "/template"),
            ("File Komponen React", "src/pages/TemplateJadwal.jsx"),
            ("Backend API Endpoint", "GET /api/template-jadwal, POST /api/template-jadwal, PUT /api/template-jadwal/:id"),
            ("Koleksi MongoDB", "template_jadwal_transportasi, master_kendaraan, master_route_nopen"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul ini menyimpan pola cetak biru (master template) jadwal keberangkatan armada harian. Template ini menentukan jam berangkat acuan, jam tiba acuan, shift (Pagi/Siang/Malam), dan kendaraan acuan untuk setiap rute sebelum digenerate menjadi jadwal harian aktual."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Master Template Builder & Generator\n"
             "- Tabel Template Jadwal per Rute Trayek Logistik\n"
             "- Form Pengaturan Jam Berangkat Standar & Jam Tiba Standar\n"
             "- Tombol Action 'Generate Jadwal Harian' untuk meng-copy template ke Jadwal Transportasi Aktif\n"
             "- Filter Shift Operasional (PAGI, SIANG, MALAM)")
        ]
    },

    # PAGE 08: TRANSPORT SCHEDULE
    {
        "filename": "08_Transport_Schedule",
        "title": "Halaman 8 — Transport Schedule (Jadwal Transportasi Harian)",
        "metadata": [
            ("Nama Modul", "Jadwal Transportasi Harian (/jadwal)"),
            ("Kategori Menu", "LOGISTICS"),
            ("Route URL", "/jadwal"),
            ("File Komponen React", "src/pages/JadwalTransportasi.jsx"),
            ("Backend API Endpoint", "GET /api/jadwal, POST /api/jadwal, PUT /api/jadwal/:id"),
            ("Koleksi MongoDB", "jadwal_transportasi, master_kendaraan, master_route_nopen"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul Jadwal Transportasi mengelola eksekusi jadwal perjalanan armada aktual pada tanggal tertentu. Berbeda dengan template, jadwal transportasi memuat tanggal berangkat riil, nomor polisi kendaraan riil yang ditugaskan, jam berangkat riil, serta status keberangkatan (READY, DEPARTED, ARRIVED, CANCELLED)."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Kalender & Monitoring Jadwal Harian\n"
             "- Date Filter Tanggal Operasional Jadwal\n"
             "- Tabel Monitoring Jam Berangkat vs Jam Realisasi Armada\n"
             "- Badge Status Keterlambatan / On-Time Departure\n"
             "- Action Button Penugasan Driver & Mobil pada Jadwal Aktif")
        ]
    },

    # PAGE 09: MILK RUN TELEMETRY & EXECUTION
    {
        "filename": "09_Milk_Run_Telemetry",
        "title": "Halaman 9 — Milk Run Telemetry & Execution",
        "metadata": [
            ("Nama Modul", "Milk Run Telemetry & Execution (/route-journey)"),
            ("Kategori Menu", "LOGISTICS"),
            ("Route URL", "/route-journey"),
            ("File Komponen React", "src/pages/RouteJourney.jsx"),
            ("Backend API Endpoint", "GET /api/route-journeys/active, POST /api/route-journeys/start, POST /api/route-journeys/:journeyId/process-stop/:seq"),
            ("Koleksi MongoDB", "route_journeys, detail_route, transaksi, master_kantor"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul Milk Run Telemetry merupakan jantung dari eksekusi perjalanan armada multi-stop. Modul ini menangani simulasi dan eksekusi ACID transaction penambahan/penurunan muatan di setiap perhentian kantor pos (Stop 1 sampai Stop N) secara real-time dengan proteksi Idempotency Key."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Visual Execution Panel & ACID Control\n"
             "- Stepper Progress Waypoint Aktif (Origin -> Waypoint Transit -> Destination Final)\n"
             "- Panel Input Scan / Load Paket per Stop Perhentian\n"
             "- Gauge Kapasitas Kendaraan Real-Time (% Muatan & Kg Sisa)\n"
             "- Tombol Eksekusi 'Proses Stop' dengan ACID Transaction MongoDB\n"
             "- Proteksi Header Idempotency Key (Mencegah pemrosesan ganda saat gangguan koneksi)")
        ]
    },

    # PAGE 10: ESTIMASI MILK RUN
    {
        "filename": "10_Estimasi_Milk_Run",
        "title": "Halaman 10 — Estimasi Milk Run (Travel Time & Route Planner)",
        "metadata": [
            ("Nama Modul", "Estimasi Milk Run (/estimasi)"),
            ("Kategori Menu", "LOGISTICS"),
            ("Route URL", "/estimasi"),
            ("File Komponen React", "src/pages/EstimasiMilkRun.jsx"),
            ("Backend API Endpoint", "GET /api/routes/:routeId/estimate"),
            ("Koleksi MongoDB", "detail_route, master_route_nopen, master_kantor"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Halaman Estimasi Milk Run berfungsi sebagai kalkulator dan perencana estimasi waktu tempuh (ETA), akumulasi jarak kilometer, serta proyeksi konsumsi kapasitas armada di setiap titik transit sebelum armada diberangkatkan."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Kalkulator Estimasi & Simulasi Muatan\n"
             "- Selector Rute Logistik & Pilihan Jenis Kendaraan Armada\n"
             "- Tabel Proyeksi ETA Kedatangan per Stop Perhentian\n"
             "- Akumulasi Total Jarak (KM) dan Total Waktu Tempuh (Jam/Menit)\n"
             "- Simulasi Beban Muatan Paket (Kg) vs Ambang Batas Maksimal Kendaraan")
        ]
    },

    # PAGE 11: GATE MONITORING
    {
        "filename": "11_Gate_Monitoring",
        "title": "Halaman 11 — Gate Monitoring (Transit Gate Terminal Hub)",
        "metadata": [
            ("Nama Modul", "Gate Monitoring Transit (/transit-monitoring)"),
            ("Kategori Menu", "LOGISTICS"),
            ("Route URL", "/transit-monitoring"),
            ("File Komponen React", "src/pages/GateMonitoring.jsx"),
            ("Backend API Endpoint", "GET /api/gate/status, POST /api/gate/scan-in, POST /api/gate/scan-out"),
            ("Koleksi MongoDB", "tracking_events, route_journeys, transaksi"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul Gate Monitoring mengawasi arus masuk (Scan In / Inbound) dan arus keluar (Scan Out / Outbound) armada logistik serta kantong pos di Pintu Gerbang Terminal Hub SPP / KCU."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Terminal Gate Radar & Inbound/Outbound Stream\n"
             "- Board Status Pintu Gerbang (Gate 1, Gate 2, Gate 3 - INBOUND / OUTBOUND / IDLE)\n"
             "- Stream Feed Antrean Kendaraan Tiba di Terminal Hub SPP\n"
             "- Form Quick Scan Barcode Kantong Pos / Resi\n"
             "- Log Verifikasi Jumlah Kantong Dibongkar vs Jumlah Manifest")
        ]
    },

    # PAGE 12: ANALITIK & LAPORAN
    {
        "filename": "12_Analitik_dan_Laporan",
        "title": "Halaman 12 — Analitik & Laporan Operasional",
        "metadata": [
            ("Nama Modul", "Analitik & Laporan Operasional (/analytics)"),
            ("Kategori Menu", "LOGISTICS / REPORTING"),
            ("Route URL", "/analytics"),
            ("File Komponen React", "src/pages/AnalyticsReport.jsx"),
            ("Backend API Endpoint", "GET /api/analytics/performance, GET /api/analytics/export-pdf"),
            ("Koleksi MongoDB", "transaksi, route_journeys, master_kendaraan"),
            ("Hak Akses Role", "Operator & Super Admin")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul Analitik & Laporan menyajikan analisis komprehensif mengenai tingkat keberhasilan pengiriman (Delivery Success Rate), efisiensi muatan armada (% Utilisasi Tonase), tren volume harian/bulanan, serta fitur pencetakan laporan resmi."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Executive Reports & Analytics Charts\n"
             "- Grafik Tren Volume Pengiriman Kiriman Pos per Rentang Tanggal\n"
             "- Laporan Utilisasi Kapasitas per Armada Kendaraan\n"
             "- Rekapitulasi Rute Paling Padat (Top Traffic Routes)\n"
             "- Ekspor Laporan Resmi ke Format PDF & Excel / CSV")
        ]
    },

    # PAGE 13: DATABASE VIEWER / TRANSAKSI
    {
        "filename": "13_Database_Viewer_Transaksi",
        "title": "Halaman 13 — Database Viewer (Tabel Transaksi Resi)",
        "metadata": [
            ("Nama Modul", "Database Viewer Transaksi (/transaksi)"),
            ("Kategori Menu", "SYSTEM (RESTRICTED)"),
            ("Route URL", "/transaksi"),
            ("File Komponen React", "src/pages/Transaksi.jsx"),
            ("Backend API Endpoint", "GET /api/transaksi?page=1&limit=25, GET /api/transaksi/stats"),
            ("Koleksi MongoDB", "transaksi, master_kantor"),
            ("Hak Akses Role", "Restricted / Super Admin & Authorized Operators")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Database Viewer merupakan modul penjelajah data transaksi mentah (raw connote documents) di database MongoDB `transaksi`. Modul ini mendukung server-side pagination, multi-field search, dan filter spesifik (State, Service, Destination Nopen/KPRK/Regional)."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Server-Side Paginated Table & Multi-Filter\n"
             "- Tabel Transaksi Resi dengan Paginasi Server (25 / 50 / 100 / 200 baris per halaman)\n"
             "- Filter Berdasarkan Status (ENTRY, LOADED, IN_TRANSIT, DELIVERED)\n"
             "- Filter Berdasarkan Jenis Layanan (Pos Express, Pos Reguler, Pos Nextday)\n"
             "- Filter Berdasarkan Nopen Tujuan / KPRK / Regional (Reg 1 - Reg 6)\n"
             "- Modal Detail JSON Raw Document Transaksi")
        ]
    },

    # PAGE 14: SYSTEM SETTINGS & COMPASS
    {
        "filename": "14_System_Settings_Compass",
        "title": "Halaman 14 — System Settings & Compass Connection",
        "metadata": [
            ("Nama Modul", "System Settings & Compass Manager (/settings & /compass)"),
            ("Kategori Menu", "SYSTEM (RESTRICTED)"),
            ("Route URL", "/settings, /compass"),
            ("File Komponen React", "src/pages/SettingsPage.jsx, src/pages/Compass.jsx"),
            ("Backend API Endpoint", "GET /api/compass/connections, POST /api/compass/connect, POST /api/compass/switch"),
            ("Koleksi MongoDB", "Konfigurasi MongoDB Multi-Server (`connections.json`)"),
            ("Hak Akses Role", "SUPER_ADMIN Only")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Modul System Settings & Compass Manager digunakan oleh Super Admin untuk mengelola koneksi database MongoDB multi-server (Primary Remote Mongo 192.168.5.219 vs Local Mongo 127.0.0.1), menguji ping koneksi, dan mengkonfigurasi parameter sistem."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### Multi-Server MongoDB Switcher & Diagnostic\n"
             "- Card Manager Koneksi MongoDB (Primary Server vs Local Fallback)\n"
             "- Indicator Status Live Connection (Badge Hijau Connected / Merah Disconnected)\n"
             "- Tombol Action 'Ping Connection' & 'Switch Connection'\n"
             "- Input Form Tambah Profil Server MongoDB Baru (URI & Database Name)")
        ]
    },

    # PAGE 15: USER PROFILE
    {
        "filename": "15_User_Profile",
        "title": "Halaman 15 — User Profile & Account Management",
        "metadata": [
            ("Nama Modul", "User Profile & Account (/profile)"),
            ("Kategori Menu", "ACCOUNT"),
            ("Route URL", "/profile"),
            ("File Komponen React", "src/pages/Profile.jsx"),
            ("Backend API Endpoint", "GET /api/auth/me, PUT /api/auth/update-profile"),
            ("Koleksi MongoDB", "users"),
            ("Hak Akses Role", "Semua Role (Operator & Super Admin)")
        ],
        "sections": [
            ("1. Deskripsi & Tujuan Utama Modul",
             "Halaman Profile menampilkan informasi akun pengguna yang sedang login (Nama Lengkap, Username, Role, NIP / ID Pegawai Pos, Kantor Operasional), serta menyediakan fitur ubah password dan manajemen sesi login."),

            ("2. Fitur-Fitur Utama & Komponen UI",
             "### User Profile Card & Security Settings\n"
             "- Kartu Informasi Pengguna dengan Badge Hak Akses (SUPER_ADMIN / OPERATOR)\n"
             "- Form Update Informasi Profil & Alamat Email\n"
             "- Form Ganti Password Sesi Login\n"
             "- Riwayat Sesi Aktivitas Login User")
        ]
    }
]

# ==============================================================================
# MAIN SCRIPT EXECUTION
# ==============================================================================

if __name__ == '__main__':
    print("[START] Starting documentation generation for all 15 IPOS5 pages...")
    print(f"[PATH] Destination Folder: {OUTPUT_DIR}\n")

    for page_info in PAGES_DATA:
        create_page_documentation(
            filename_base=page_info["filename"],
            title=page_info["title"],
            metadata=page_info["metadata"],
            sections=page_info["sections"]
        )

    print("\n[SUCCESS] ALL 16 MARKDOWN (.md) & 16 WORD DOCUMENT (.docx) FILES GENERATED SUCCESSFULLY!")
