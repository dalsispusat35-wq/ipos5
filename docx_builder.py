import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

OUTPUT_DIR = r"c:\Users\Asus\Documents\POSIND\IPOS5\ipos5\dokumentasi_halaman_ipos5"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Define Helper for Word Styling
def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_docx_from_md_data(title, metadata, sections, output_path):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"DOKUMENTASI TEKNIS & OPERASIONAL\n{title.upper()}")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(13, 27, 56) # Deep Navy
    title_p.paragraph_format.space_after = Pt(12)

    # Subtitle / System Tag
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("IPOS5 Routing & Schedule Management System — PT Pos Indonesia (Persero)")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(10.5)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(232, 67, 31) # Pos Orange
    sub_p.paragraph_format.space_after = Pt(18)

    # Metadata Table
    meta_table = doc.add_table(rows=len(metadata), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    for idx, (k, v) in enumerate(metadata):
        row = meta_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.5)
        
        set_cell_background(cell_k, "0D1B38") # Deep Navy
        set_cell_background(cell_v, "F1F5F9") # Slate Light
        
        set_cell_margins(cell_k, top=120, bottom=120, left=150, right=150)
        set_cell_margins(cell_v, top=120, bottom=120, left=150, right=150)

        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.name = 'Arial'
        rk.font.size = Pt(9.5)
        rk.font.bold = True
        rk.font.color.rgb = RGBColor(255, 255, 255)

        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.name = 'Arial'
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Dynamic Sections
    for sec_title, sec_content in sections:
        # Heading 2
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(sec_title)
        h_run.font.name = 'Arial'
        h_run.font.size = Pt(13.5)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(2, 132, 199) # Sky Blue
        h_p.paragraph_format.space_before = Pt(16)
        h_p.paragraph_format.space_after = Pt(6)

        # Content lines
        for line in sec_content.split('\n'):
            line_str = line.strip()
            if not line_str:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            if line_str.startswith('- ') or line_str.startswith('* '):
                p.paragraph_format.left_indent = Inches(0.25)
                r_bullet = p.add_run("• ")
                r_bullet.font.name = 'Arial'
                r_bullet.font.bold = True
                r_bullet.font.color.rgb = RGBColor(232, 67, 31)
                
                r_text = p.add_run(line_str[2:])
                r_text.font.name = 'Arial'
                r_text.font.size = Pt(10)
                r_text.font.color.rgb = RGBColor(51, 65, 85)
            elif line_str.startswith('### '):
                r_sub = p.add_run(line_str[4:])
                r_sub.font.name = 'Arial'
                r_sub.font.size = Pt(11)
                r_sub.font.bold = True
                r_sub.font.color.rgb = RGBColor(13, 27, 56)
                p.paragraph_format.space_before = Pt(10)
            elif line_str.startswith('⚠️') or line_str.startswith('📌') or line_str.startswith('💡'):
                set_cell_background(doc.add_table(rows=1, cols=1).rows[0].cells[0], "FEF3C7")
                # simplified paragraph
                r_box = p.add_run(line_str)
                r_box.font.name = 'Arial'
                r_box.font.size = Pt(9.5)
                r_box.font.bold = True
                r_box.font.color.rgb = RGBColor(180, 83, 9)
            else:
                r_text = p.add_run(line_str)
                r_text.font.name = 'Arial'
                r_text.font.size = Pt(10)
                r_text.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(output_path)
    print(f"Generated DOCX: {os.path.basename(output_path)}")

print("DOCX builder ready.")
