import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from generate_all_page_docs import PAGES_DATA, OUTPUT_DIR, set_cell_background, set_cell_margins

def create_master_combined_docx():
    master_docx_path = os.path.join(OUTPUT_DIR, "DOKUMENTASI_LENGKAP_ALL_PAGES_IPOS5.docx")
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # Cover Header
    tp = doc.add_paragraph()
    tr = tp.add_run("BUNTAI DOKUMENTASI LENGKAP SELURUH HALAMAN APLIKASI\nIPOS5 ROUTING & SCHEDULE MANAGEMENT")
    tr.font.name = 'Arial'
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(13, 27, 56) # Deep Navy
    tp.paragraph_format.space_after = Pt(6)

    sub = doc.add_paragraph()
    sr = sub.add_run("PT POS INDONESIA (PERSERO) — INTEGRATED POSTAL OPERATIONAL SYSTEM REDESIGN V2.5")
    sr.font.name = 'Arial'
    sr.font.size = Pt(11)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor(232, 67, 31) # Pos Orange
    sub.paragraph_format.space_after = Pt(24)

    # Master Table of Contents Banner
    h0 = doc.add_paragraph()
    h0r = h0.add_run("1. RINGKASAN DAFTAR MODUL HALAMAN")
    h0r.font.name = 'Arial'
    h0r.font.size = Pt(14)
    h0r.font.bold = True
    h0r.font.color.rgb = RGBColor(2, 132, 199)
    h0.paragraph_format.space_after = Pt(10)

    # Summary List Table
    summary_table = doc.add_table(rows=len(PAGES_DATA), cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.autofit = False

    for idx, page in enumerate(PAGES_DATA):
        row = summary_table.rows[idx]
        c1, c2, c3 = row.cells[0], row.cells[1], row.cells[2]
        c1.width = Inches(0.8)
        c2.width = Inches(3.2)
        c3.width = Inches(2.7)
        
        bg_col = "0D1B38" if idx == 0 else ("F1F5F9" if idx % 2 == 1 else "FFFFFF")
        text_col = RGBColor(255, 255, 255) if idx == 0 else RGBColor(15, 23, 42)

        set_cell_background(c1, bg_col)
        set_cell_background(c2, bg_col)
        set_cell_background(c3, bg_col)
        
        set_cell_margins(c1)
        set_cell_margins(c2)
        set_cell_margins(c3)

        # Route Metadata URL
        route_val = "-"
        for k, v in page["metadata"]:
            if k == "Route URL":
                route_val = v
                break

        r1 = c1.paragraphs[0].add_run(f"#{idx}")
        r1.font.bold = True
        r1.font.color.rgb = text_col
        
        r2 = c2.paragraphs[0].add_run(page["title"])
        r2.font.bold = True
        r2.font.color.rgb = text_col
        
        r3 = c3.paragraphs[0].add_run(f"Route: {route_val}")
        r3.font.color.rgb = text_col

    doc.add_page_break()

    # Iterate & Add Each Page Documentation
    for p_idx, page in enumerate(PAGES_DATA):
        if p_idx == 0:
            continue # index page already summarized

        # Title
        hp = doc.add_paragraph()
        hr = hp.add_run(f"MODUL HALAMAN {p_idx} — {page['title'].upper()}")
        hr.font.name = 'Arial'
        hr.font.size = Pt(16)
        hr.font.bold = True
        hr.font.color.rgb = RGBColor(13, 27, 56)
        hp.paragraph_format.space_before = Pt(16)
        hp.paragraph_format.space_after = Pt(8)

        # Metadata Table
        meta_table = doc.add_table(rows=len(page["metadata"]), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.autofit = False

        for idx, (k, v) in enumerate(page["metadata"]):
            row = meta_table.rows[idx]
            ck, cv = row.cells[0], row.cells[1]
            ck.width = Inches(2.2)
            cv.width = Inches(4.5)
            set_cell_background(ck, "0D1B38")
            set_cell_background(cv, "F8FAFC")
            set_cell_margins(ck)
            set_cell_margins(cv)

            pk = ck.paragraphs[0]
            rk = pk.add_run(k)
            rk.font.name = 'Arial'
            rk.font.size = Pt(9.5)
            rk.font.bold = True
            rk.font.color.rgb = RGBColor(255, 255, 255)

            pv = cv.paragraphs[0]
            rv = pv.add_run(v)
            rv.font.name = 'Arial'
            rv.font.size = Pt(9.5)
            rv.font.color.rgb = RGBColor(15, 23, 42)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # Sections
        for sec_title, sec_content in page["sections"]:
            sec_p = doc.add_paragraph()
            sec_r = sec_p.add_run(sec_title)
            sec_r.font.name = 'Arial'
            sec_r.font.size = Pt(12)
            sec_r.font.bold = True
            sec_r.font.color.rgb = RGBColor(2, 132, 199)
            sec_p.paragraph_format.space_before = Pt(12)
            sec_p.paragraph_format.space_after = Pt(4)

            for line in sec_content.split('\n'):
                line_s = line.strip()
                if not line_s:
                    continue

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15

                if line_s.startswith('- ') or line_s.startswith('* '):
                    p.paragraph_format.left_indent = Inches(0.25)
                    rb = p.add_run("• ")
                    rb.font.bold = True
                    rb.font.color.rgb = RGBColor(232, 67, 31)

                    rt = p.add_run(line_s[2:])
                    rt.font.name = 'Arial'
                    rt.font.size = Pt(9.5)
                    rt.font.color.rgb = RGBColor(51, 65, 85)
                elif line_s.startswith('### '):
                    rs = p.add_run(line_s[4:])
                    rs.font.name = 'Arial'
                    rs.font.size = Pt(10)
                    rs.font.bold = True
                    rs.font.color.rgb = RGBColor(13, 27, 56)
                    p.paragraph_format.space_before = Pt(6)
                else:
                    rt = p.add_run(line_s)
                    rt.font.name = 'Arial'
                    rt.font.size = Pt(9.5)
                    rt.font.color.rgb = RGBColor(51, 65, 85)

        doc.add_page_break()

    doc.save(master_docx_path)
    print(f"[SUCCESS] Master Combined DOCX document saved to: {master_docx_path}")

if __name__ == '__main__':
    create_master_combined_docx()
